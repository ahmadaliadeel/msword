"""Unit tests for `msword.io.docx_import`.

The fixture is built **programmatically** with `python-docx` so we don't have
to commit a binary `.docx` to the repo (per the unit brief). It contains:

* 1 heading
* 2 plain paragraphs
* 1 bullet list with 3 items
* 1 inline 1x1 PNG image
* 1 2x2 table

We then assert exact block counts and that exactly 1 asset was registered.
"""

from __future__ import annotations

import logging
import struct
import zlib
from pathlib import Path

import docx
import pytest

from msword.io.docx_import import (
    Document,
    HeadingBlock,
    ImageBlock,
    ListBlock,
    ParagraphBlock,
    TableBlock,
    import_docx,
)

# ---------------------------------------------------------------------------
# Tiny 1x1 PNG, generated in-memory so no binary fixture is committed.
# Reference: PNG spec, IHDR + IDAT + IEND minimal valid file.
# ---------------------------------------------------------------------------


def _make_1x1_png() -> bytes:
    sig = b"\x89PNG\r\n\x1a\n"

    def _chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    # 1x1, 8-bit grayscale.
    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 0, 0, 0, 0)
    # 1 scanline: filter byte 0 + 1 sample byte.
    idat = zlib.compress(b"\x00\x00")
    return sig + _chunk(b"IHDR", ihdr) + _chunk(b"IDAT", idat) + _chunk(b"IEND", b"")


@pytest.fixture()
def sample_docx(tmp_path: Path) -> Path:
    """Build the sample DOCX described in the test file's module docstring."""

    doc = docx.Document()

    doc.add_heading("Hello world", level=1)
    doc.add_paragraph("First plain paragraph with some body text.")
    doc.add_paragraph("Second plain paragraph for good measure.")

    # Bullet list — `python-docx` ships a built-in "List Bullet" style.
    for item in ("alpha", "beta", "gamma"):
        doc.add_paragraph(item, style="List Bullet")

    # Inline image — write the PNG bytes to a temp file `python-docx` can open.
    png_path = tmp_path / "px.png"
    png_path.write_bytes(_make_1x1_png())
    doc.add_picture(str(png_path))

    # 2x2 table.
    table = doc.add_table(rows=2, cols=2)
    for r, row in enumerate(table.rows):
        for c, cell in enumerate(row.cells):
            cell.text = f"r{r}c{c}"

    out = tmp_path / "sample.docx"
    doc.save(str(out))
    return out


def test_import_docx_returns_document(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    assert isinstance(out, Document)


def test_block_counts_match_fixture(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    [story] = out.stories.values()

    headings = [b for b in story.blocks if isinstance(b, HeadingBlock)]
    paragraphs = [b for b in story.blocks if isinstance(b, ParagraphBlock)]
    lists = [b for b in story.blocks if isinstance(b, ListBlock)]
    images = [b for b in story.blocks if isinstance(b, ImageBlock)]
    tables = [b for b in story.blocks if isinstance(b, TableBlock)]

    assert len(headings) == 1
    # `python-docx` `add_picture` inserts a paragraph *just for* the image; we
    # collapse that empty-text paragraph into the ImageBlock, so the only
    # plain paragraphs left are the two we wrote explicitly.
    assert len(paragraphs) == 2
    assert len(lists) == 1
    assert len(lists[0].items) == 3
    assert len(images) == 1
    assert len(tables) == 1
    assert tables[0].rows == 2 and tables[0].cols == 2


def test_one_asset_registered(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    assert len(out.assets) == 1
    asset = next(iter(out.assets.values()))
    # Content-addressed by SHA-256 (spec §4).
    assert len(asset.sha256) == 64
    assert asset.data.startswith(b"\x89PNG\r\n\x1a\n")
    assert asset.content_type.startswith("image/")


def test_master_page_geometry(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    [master] = out.master_pages.values()
    assert master.name == "A-Master"
    assert master.page_width > 0
    assert master.page_height > 0
    assert len(master.frames) == 1
    frame = master.frames[0]
    # The auto-flowed frame's story_ref points at the imported story.
    [story] = out.stories.values()
    assert frame.story_ref == story.id


def test_paragraph_styles_imported(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    # python-docx's blank document carries the standard built-in styles, so
    # at minimum the heading style we used must be present.
    assert "Heading 1" in out.paragraph_styles


def test_heading_level(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    [story] = out.stories.values()
    headings = [b for b in story.blocks if isinstance(b, HeadingBlock)]
    assert headings[0].level == 1
    assert headings[0].runs[0].text == "Hello world"


def test_lossy_features_warn(tmp_path: Path, caplog: pytest.LogCaptureFixture) -> None:
    """Headers/footers must trigger a `logging.warning` (spec mapping rules)."""

    doc = docx.Document()
    doc.add_paragraph("body")
    section = doc.sections[0]
    # Force a non-empty header and explicitly mark not-linked-to-previous.
    section.header.is_linked_to_previous = False
    section.header.add_paragraph("running title")

    out_path = tmp_path / "with_header.docx"
    doc.save(str(out_path))

    with caplog.at_level(logging.WARNING, logger="msword.io.docx_import"):
        import_docx(out_path)
    assert any("header" in rec.message.lower() for rec in caplog.records)


def test_missing_file_raises(tmp_path: Path) -> None:
    with pytest.raises(FileNotFoundError):
        import_docx(tmp_path / "does-not-exist.docx")


def test_io_module_has_no_qt_import() -> None:
    """`io/` is pure I/O per the anchor invariant (spec §3 / unit brief)."""

    import sys

    import msword.io.docx_import

    qt_modules = [name for name in sys.modules if name.startswith("PySide6")]
    # Importing the module on its own must not have pulled in Qt.
    # (Other tests may have, so we check that the module itself doesn't
    # reference any Qt names in its globals.)
    mod_globals = vars(msword.io.docx_import)
    for v in mod_globals.values():
        mod_name = getattr(v, "__module__", "") or ""
        assert not mod_name.startswith("PySide6"), (
            f"io.docx_import leaks Qt symbol from {mod_name}"
        )
    # Suppress unused-var warning when Qt was never loaded:
    _ = qt_modules


def test_run_text_round_trip(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    [story] = out.stories.values()
    paragraphs = [b for b in story.blocks if isinstance(b, ParagraphBlock)]
    body_text = " ".join(r.text for p in paragraphs for r in p.runs)
    assert "First plain paragraph" in body_text
    assert "Second plain paragraph" in body_text


def test_table_cells_have_content(sample_docx: Path) -> None:
    out = import_docx(sample_docx)
    [story] = out.stories.values()
    [table] = (b for b in story.blocks if isinstance(b, TableBlock))
    flat = [
        run.text
        for row in table.cells
        for cell in row
        for blk in cell
        if isinstance(blk, ParagraphBlock)
        for run in blk.runs
    ]
    assert "r0c0" in flat
    assert "r1c1" in flat


def test_pathlike_argument_accepted(sample_docx: Path) -> None:
    # Sanity check that `os.PathLike` works (we type-hinted `import_docx`).
    out = import_docx(sample_docx)  # `Path` is `os.PathLike[str]`.
    assert out is not None
