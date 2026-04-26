"""DOCX importer — reads a `.docx` file and produces a `Document` model tree.

This module owns unit #11 of the spec (`io-docx-import`). It is **pure I/O**:
no Qt, no rendering. The full model package will land via sibling units; until
then this module ships local stubs for `Document`, `Page`, `MasterPage`,
`TextFrame`, `Story`, the block hierarchy, `Run`, the style classes, and
`Asset` so that import + tests work in isolation per spec §12.1 ("Units that
need a dependency stub it locally").

Mapping rules (spec §8 "Interop / DOCX import"):

* The DOCX is collapsed onto a **single A-Master** built from the first
  section's geometry; one auto-flowed `TextFrame` covers the live area.
* Word paragraph/character styles → `ParagraphStyle` / `CharacterStyle`
  (1:1 by name; "based-on" is preserved when present).
* `Heading N` paragraphs → `HeadingBlock(level=N)` (Title → level 0).
* List paragraphs (paragraphs carrying `w:numPr`) → `ListBlock` items;
  consecutive list paragraphs collapse into a single `ListBlock`.
* Inline images (a `w:drawing` inside a run) → `ImageBlock` + `Asset`,
  content-addressed by SHA-256 (spec §4 / §8).
* Tables → `TableBlock` (cells hold a list of paragraph blocks).
* Everything else → `ParagraphBlock`.

Lossy features (text boxes, headers/footers, footnotes, comments,
endnotes) are not represented in the imported model and emit a
`logging.warning` so the user is aware (spec §15).
"""

from __future__ import annotations

import hashlib
import logging
import os
import uuid
from dataclasses import dataclass, field
from pathlib import Path
from typing import TYPE_CHECKING, Any

import docx
from docx.oxml.ns import qn

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument
    from docx.table import Table as DocxTable
    from docx.text.paragraph import Paragraph as DocxParagraph
    from docx.text.run import Run as DocxRun

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Local model stubs (§12.1: stub until owning unit lands).
# These mirror the data shapes the spec mandates in §4 and §4.2 so that
# downstream code can already program against the public attribute names.
# ---------------------------------------------------------------------------


def _new_id(prefix: str) -> str:
    return f"{prefix}_{uuid.uuid4().hex[:8]}"


@dataclass
class Asset:
    """Content-addressed binary asset (image, font). Spec §4."""

    sha256: str
    content_type: str
    data: bytes
    filename: str | None = None


@dataclass
class Run:
    """Inline styling unit. Spec §4.2 lists the full mark set; we ship the
    subset we can actually recover from a DOCX."""

    text: str = ""
    bold: bool | None = None
    italic: bool | None = None
    underline: bool | None = None
    font_ref: str | None = None
    size: float | None = None
    character_style_ref: str | None = None


@dataclass
class Block:
    """Base block. Spec §4.2."""

    id: str = field(default_factory=lambda: _new_id("blk"))


@dataclass
class ParagraphBlock(Block):
    runs: list[Run] = field(default_factory=list)
    paragraph_style_ref: str | None = None


@dataclass
class HeadingBlock(Block):
    level: int = 1
    runs: list[Run] = field(default_factory=list)
    paragraph_style_ref: str | None = None


@dataclass
class ListItem:
    blocks: list[Block] = field(default_factory=list)
    level: int = 0


@dataclass
class ListBlock(Block):
    kind: str = "bullet"  # "bullet" | "ordered" | "todo"
    items: list[ListItem] = field(default_factory=list)


@dataclass
class ImageBlock(Block):
    asset_ref: str = ""
    layout: str = "inline"  # "inline" | "float" | "full-width"
    caption: str | None = None


@dataclass
class TableBlock(Block):
    rows: int = 0
    cols: int = 0
    cells: list[list[list[Block]]] = field(default_factory=list)


@dataclass
class Story:
    """A story is a flow of blocks. Spec §4.2."""

    id: str = field(default_factory=lambda: _new_id("story"))
    blocks: list[Block] = field(default_factory=list)


@dataclass
class TextFrame:
    """Frame holding a story. Spec §4.1."""

    id: str = field(default_factory=lambda: _new_id("frame"))
    page_id: str = ""
    x: float = 0.0
    y: float = 0.0
    w: float = 0.0
    h: float = 0.0
    columns: int = 1
    gutter: float = 0.0
    text_direction: str = "inherit"
    story_ref: str = ""
    story_index: int = 0


@dataclass
class MasterPage:
    """Template page. Spec §4."""

    id: str = field(default_factory=lambda: _new_id("master"))
    name: str = "A-Master"
    page_width: float = 0.0
    page_height: float = 0.0
    margin_top: float = 0.0
    margin_bottom: float = 0.0
    margin_left: float = 0.0
    margin_right: float = 0.0
    frames: list[TextFrame] = field(default_factory=list)


@dataclass
class Page:
    """A page references a master and may override its items. Spec §4."""

    id: str = field(default_factory=lambda: _new_id("page"))
    master_ref: str = ""


@dataclass
class ParagraphStyle:
    name: str
    based_on: str | None = None
    properties: dict[str, Any] = field(default_factory=dict)


@dataclass
class CharacterStyle:
    name: str
    based_on: str | None = None
    properties: dict[str, Any] = field(default_factory=dict)


@dataclass
class Document:
    """Top-level model. Spec §4."""

    meta: dict[str, Any] = field(default_factory=dict)
    paragraph_styles: dict[str, ParagraphStyle] = field(default_factory=dict)
    character_styles: dict[str, CharacterStyle] = field(default_factory=dict)
    master_pages: dict[str, MasterPage] = field(default_factory=dict)
    pages: list[Page] = field(default_factory=list)
    stories: dict[str, Story] = field(default_factory=dict)
    assets: dict[str, Asset] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Importer
# ---------------------------------------------------------------------------


# Word's default page-size unit is EMU (English Metric Unit); 914400 per inch.
# We keep raw EMU values: a downstream layout unit converts to its own units.
_EMU_PER_PT = 12700  # 914400 / 72


def _emu_to_pt(value: int | None) -> float:
    if value is None:
        return 0.0
    return float(value) / _EMU_PER_PT


def import_docx(path: str | os.PathLike[str]) -> Document:
    """Open `path` with `python-docx` and return a populated `Document`.

    See module docstring for the mapping rules.
    """

    src = Path(path)
    if not src.exists():
        raise FileNotFoundError(f"DOCX not found: {src}")

    docx_doc: DocxDocument = docx.Document(str(src))
    out = Document()
    out.meta = {"source_path": str(src), "source_kind": "docx"}

    _warn_lossy(docx_doc)
    _import_styles(docx_doc, out)
    master = _build_master(docx_doc, out)
    page = Page(master_ref=master.id)
    out.pages.append(page)

    story = Story()
    out.stories[story.id] = story
    if master.frames:
        master.frames[0].story_ref = story.id
        master.frames[0].page_id = page.id

    _import_body(docx_doc, story, out)

    return out


# -- styles ------------------------------------------------------------------


def _import_styles(docx_doc: DocxDocument, out: Document) -> None:
    """Map Word styles to `ParagraphStyle` / `CharacterStyle` (§8 interop)."""

    try:
        styles = docx_doc.styles
    except Exception:  # pragma: no cover - defensive
        log.warning("DOCX has no styles element; using defaults")
        return

    for s in styles:
        # `s.type` is a WD_STYLE_TYPE enum; values 1==paragraph, 2==character.
        try:
            stype = s.type
        except Exception:
            continue
        name = getattr(s, "name", None)
        if not name:
            continue
        based_on_obj = getattr(s, "base_style", None)
        based_on = based_on_obj.name if based_on_obj is not None else None
        # WD_STYLE_TYPE.PARAGRAPH == 1, CHARACTER == 2; we don't import table /
        # list styles in this unit (separate model concerns).
        if int(stype) == 1:
            out.paragraph_styles[name] = ParagraphStyle(name=name, based_on=based_on)
        elif int(stype) == 2:
            out.character_styles[name] = CharacterStyle(name=name, based_on=based_on)


# -- master / page geometry --------------------------------------------------


def _build_master(docx_doc: DocxDocument, out: Document) -> MasterPage:
    """Build the single A-Master from the first section (spec §8)."""

    section = docx_doc.sections[0] if len(docx_doc.sections) else None

    def _val(attr: str, default: float = 0.0) -> float:
        if section is None:
            return default
        v = getattr(section, attr, None)
        return _emu_to_pt(v) if v is not None else default

    page_w = _val("page_width")
    page_h = _val("page_height")
    m_top = _val("top_margin")
    m_bot = _val("bottom_margin")
    m_left = _val("left_margin")
    m_right = _val("right_margin")

    # The auto-flowed text frame covers the live area (page minus margins).
    frame = TextFrame(
        x=m_left,
        y=m_top,
        w=max(0.0, page_w - m_left - m_right),
        h=max(0.0, page_h - m_top - m_bot),
        columns=1,
    )
    master = MasterPage(
        name="A-Master",
        page_width=page_w,
        page_height=page_h,
        margin_top=m_top,
        margin_bottom=m_bot,
        margin_left=m_left,
        margin_right=m_right,
        frames=[frame],
    )
    out.master_pages[master.id] = master
    return master


# -- body walk ---------------------------------------------------------------


def _import_body(docx_doc: DocxDocument, story: Story, out: Document) -> None:
    """Walk the document body, in order, emitting blocks into `story`."""

    pending_list: list[tuple[DocxParagraph, str, int]] = []

    def _flush_list() -> None:
        if not pending_list:
            return
        kind = pending_list[0][1]
        items: list[ListItem] = []
        for para, _kind, level in pending_list:
            pblock = _paragraph_to_block(para)
            items.append(ListItem(blocks=[pblock], level=level))
        story.blocks.append(ListBlock(kind=kind, items=items))
        pending_list.clear()

    for item in docx_doc.iter_inner_content():
        # Tables flush any pending list first.
        if _is_table(item):
            _flush_list()
            story.blocks.append(_table_to_block(item))  # type: ignore[arg-type]
            continue

        para: DocxParagraph = item  # type: ignore[assignment]
        list_info = _list_info(para)
        if list_info is not None:
            kind, level = list_info
            if pending_list and pending_list[0][1] != kind:
                _flush_list()
            pending_list.append((para, kind, level))
            continue

        _flush_list()

        heading_level = _heading_level(para)
        if heading_level is not None:
            story.blocks.append(_paragraph_to_heading(para, heading_level))
            continue

        # If the paragraph contains an inline image, emit an ImageBlock
        # (alongside any text). We emit text first as a paragraph, then the
        # image, to keep block order stable.
        image_emitted = _maybe_emit_image_blocks(para, story, out)
        if image_emitted and not _has_visible_text(para):
            continue
        story.blocks.append(_paragraph_to_block(para))

    _flush_list()


def _is_table(item: object) -> bool:
    return item.__class__.__name__ == "Table"


# -- paragraph helpers -------------------------------------------------------


def _heading_level(para: DocxParagraph) -> int | None:
    """Return Heading level 1-6 (or 0 for Title) if the style is a heading."""
    style = para.style
    if style is None:
        return None
    name = getattr(style, "name", "") or ""
    if name == "Title":
        return 0
    if name.startswith("Heading "):
        try:
            level = int(name.split(" ", 1)[1])
        except ValueError:
            return None
        if 1 <= level <= 6:
            return level
    return None


def _list_info(para: DocxParagraph) -> tuple[str, int] | None:
    """Detect list paragraphs. Returns (kind, level) or None.

    Word marks list membership in two places: directly on the paragraph
    (``w:pPr/w:numPr``) or via a paragraph style whose ``w:pPr`` carries
    ``w:numPr``. We also recognise the canonical built-in style names
    (``List Bullet``, ``List Number``, plus their ``"List Bullet 2"``
    variants) since those are what ``python-docx`` and Word emit by default.

    Bullet-vs-ordered classification is a conservative heuristic in this
    unit: the proper resolution lives in the numbering part
    (``w:abstractNumId`` → ``w:numFmt``) and will land with unit #6.
    """
    style_name = (getattr(para.style, "name", "") or "")

    # Direct paragraph-level numPr.
    numPr = None
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))

    # Style-level numPr (walk the based-on chain).
    style_has_num = False
    if numPr is None:
        style_obj = para.style
        seen: set[int] = set()
        while style_obj is not None and id(style_obj) not in seen:
            seen.add(id(style_obj))
            style_el = getattr(style_obj, "element", None)
            if style_el is not None:
                style_pPr = style_el.find(qn("w:pPr"))
                if style_pPr is not None and style_pPr.find(qn("w:numPr")) is not None:
                    style_has_num = True
                    break
            style_obj = getattr(style_obj, "base_style", None)

    # Heuristic fallback: the canonical built-in list style names.
    name_implies_list = bool(
        style_name
        and (
            style_name.startswith("List Bullet")
            or style_name.startswith("List Number")
            or style_name.startswith("List Paragraph")
        )
    )

    if numPr is None and not style_has_num and not name_implies_list:
        return None

    level = 0
    if numPr is not None:
        ilvl_el = numPr.find(qn("w:ilvl"))
        if ilvl_el is not None:
            val = ilvl_el.get(qn("w:val"))
            try:
                level = int(val) if val is not None else 0
            except ValueError:
                level = 0

    kind = "ordered" if "number" in style_name.lower() else "bullet"
    return kind, level


def _has_visible_text(para: DocxParagraph) -> bool:
    return bool(para.text and para.text.strip())


def _paragraph_to_heading(para: DocxParagraph, level: int) -> HeadingBlock:
    return HeadingBlock(
        level=level,
        runs=_runs_from_paragraph(para),
        paragraph_style_ref=_para_style_name(para),
    )


def _paragraph_to_block(para: DocxParagraph) -> ParagraphBlock:
    return ParagraphBlock(
        runs=_runs_from_paragraph(para),
        paragraph_style_ref=_para_style_name(para),
    )


def _para_style_name(para: DocxParagraph) -> str | None:
    style = para.style
    if style is None:
        return None
    return getattr(style, "name", None)


def _runs_from_paragraph(para: DocxParagraph) -> list[Run]:
    runs: list[Run] = []
    for r in para.runs:
        text = r.text or ""
        if not text:
            # Skip empty runs (often image-only or formatting placeholders).
            continue
        runs.append(_run_to_model(r))
    if not runs:
        # Preserve empty paragraphs as a single empty run for round-tripping.
        runs.append(Run(text=""))
    return runs


def _run_to_model(r: DocxRun) -> Run:
    char_style = None
    try:
        cs = r.style
        if cs is not None:
            char_style = getattr(cs, "name", None)
    except Exception:
        char_style = None
    font = r.font
    return Run(
        text=r.text or "",
        bold=r.bold,
        italic=r.italic,
        underline=bool(r.underline) if r.underline is not None else None,
        font_ref=getattr(font, "name", None),
        size=float(font.size.pt) if font.size is not None else None,
        character_style_ref=char_style,
    )


# -- image helpers -----------------------------------------------------------


def _maybe_emit_image_blocks(
    para: DocxParagraph, story: Story, out: Document
) -> bool:
    """Scan paragraph runs for `w:drawing`; emit ImageBlocks for each.

    Returns True if at least one image was emitted.
    """
    emitted = False
    for run in para.runs:
        for blip_rid in _iter_blip_rids(run):
            asset = _load_image_asset(run, blip_rid)
            if asset is None:
                continue
            out.assets[asset.sha256] = asset
            story.blocks.append(ImageBlock(asset_ref=asset.sha256, layout="inline"))
            emitted = True
    return emitted


_NS_DRAWINGML = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_RELS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_BLIP_TAG = f"{{{_NS_DRAWINGML}}}blip"
_R_EMBED_ATTR = f"{{{_NS_RELS}}}embed"


def _iter_blip_rids(run: DocxRun) -> list[str]:
    """Return the relationship IDs referenced by `<a:blip r:embed="rIdN">`
    elements inside a run's `<w:drawing>` children.
    """
    rids: list[str] = []
    for drawing in run._r.iter(qn("w:drawing")):
        for blip in drawing.iter(_BLIP_TAG):
            rid = blip.get(_R_EMBED_ATTR)
            if rid:
                rids.append(rid)
    return rids


def _load_image_asset(run: DocxRun, rid: str) -> Asset | None:
    try:
        image_part = run.part.related_parts[rid]
    except KeyError:
        log.warning("Image rId %s not found in document part relationships", rid)
        return None
    blob: bytes = image_part.blob
    sha = hashlib.sha256(blob).hexdigest()
    content_type = getattr(image_part, "content_type", "application/octet-stream")
    filename = None
    partname = getattr(image_part, "partname", None)
    if partname is not None:
        filename = str(partname).rsplit("/", 1)[-1]
    return Asset(sha256=sha, content_type=content_type, data=blob, filename=filename)


# -- table helpers -----------------------------------------------------------


def _table_to_block(table: DocxTable) -> TableBlock:
    rows = len(table.rows)
    cols = len(table.columns) if rows else 0
    cells: list[list[list[Block]]] = []
    for row in table.rows:
        row_cells: list[list[Block]] = []
        for cell in row.cells:
            cell_blocks: list[Block] = [
                _paragraph_to_block(para) for para in cell.paragraphs
            ]
            row_cells.append(cell_blocks)
        cells.append(row_cells)
    return TableBlock(rows=rows, cols=cols, cells=cells)


# -- lossy-feature warnings --------------------------------------------------


_LOSSY_TAGS: tuple[tuple[str, str], ...] = (
    ("w:txbxContent", "text boxes"),
    ("w:footnoteReference", "footnotes"),
    ("w:endnoteReference", "endnotes"),
    ("w:commentReference", "comments"),
)


def _warn_lossy(docx_doc: DocxDocument) -> None:
    """Emit `logging.warning` for features we don't represent yet."""

    body = docx_doc.element.body
    for tag, label in _LOSSY_TAGS:
        if body.find(f".//{qn(tag)}") is not None:
            log.warning("DOCX %s are not imported (lossy)", label)

    for kind in ("header", "footer"):
        if _any_section_has_content(docx_doc, kind):
            log.warning("DOCX section %s content is not imported (lossy)", kind)


def _any_section_has_content(docx_doc: DocxDocument, kind: str) -> bool:
    for sect in docx_doc.sections:
        try:
            part = getattr(sect, kind)
            if part.is_linked_to_previous is False and any(
                p.text.strip() for p in part.paragraphs
            ):
                return True
        except Exception:
            continue
    return False


__all__ = [
    "Asset",
    "Block",
    "CharacterStyle",
    "Document",
    "HeadingBlock",
    "ImageBlock",
    "ListBlock",
    "ListItem",
    "MasterPage",
    "Page",
    "ParagraphBlock",
    "ParagraphStyle",
    "Run",
    "Story",
    "TableBlock",
    "TextFrame",
    "import_docx",
]
